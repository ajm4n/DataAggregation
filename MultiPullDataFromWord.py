import docx
import csv
import tkinter as tk
from tkinter import filedialog

# Create a GUI window
window = tk.Tk()
window.title('Word Table to CSV Converter')

# Create a label for the paths to the Word documents
doc_label = tk.Label(window, text='Path to Word Document(s):')
doc_label.pack()

# Create a listbox to display the chosen Word documents
doc_listbox = tk.Listbox(window, width=50, height=5)
doc_listbox.pack()

# Create a button to open a file dialog to choose the Word documents
def choose_docs():
    doc_paths = filedialog.askopenfilenames(filetypes=[('Word Documents', '*.docx')])
    doc_listbox.delete(0, tk.END)
    for doc_path in doc_paths:
        doc_listbox.insert(tk.END, doc_path)

doc_button = tk.Button(window, text='Choose', command=choose_docs)
doc_button.pack()

# Create a label for the output file name
csv_label = tk.Label(window, text='Output File Name:')
csv_label.pack()

# Create an entry field for the output file name
csv_entry = tk.Entry(window, width=50)
csv_entry.pack()

# Create a button to open a file dialog to choose the output file name and convert the Word tables to CSV
def convert():
    doc_paths = doc_listbox.get(0, tk.END)
    csv_path = csv_entry.get()
    if not csv_path.endswith('.csv'):
        csv_path += '.csv'
    with open(csv_path, 'w', newline='') as f:
        writer = csv.writer(f)
        for doc_path in doc_paths:
            doc = docx.Document(doc_path)
            for table in doc.tables:
                headers = [cell.text.strip() for cell in table.row_cells(0)]
                if 'Vulnerability' in headers and 'Root Cause' in headers and 'Rating' in headers:
                    for row in table.rows:
                        writer.writerow([cell.text for cell in row.cells])
    result_label.config(text='Conversion complete!')

csv_button = tk.Button(window, text='Convert', command=convert)
csv_button.pack()

# Create a label to show the result of the conversion
result_label = tk.Label(window, text='')
result_label.pack()

# Run the GUI window
window.mainloop()
